#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate Word document: Laravel (Mindlytics) deployment guide using Git.
Output: Mindlytics_Deploy_Guide.docx in project root.
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Install first: pip install python-docx")
    sys.exit(1)

BODY_FONT = "Calibri"
HEADING_FONT = "Calibri"
CODE_FONT = "Consolas"
OUTPUT_NAME = "Mindlytics_Deploy_Guide.docx"


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = HEADING_FONT
    run.font.bold = True
    if level == 0:
        run.font.size = Pt(22)
        p.paragraph_format.space_after = Pt(6)
    elif level == 1:
        run.font.size = Pt(16)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(12)
    return p


def add_para(doc, text, bold=False, font_size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(font_size)
    run.font.bold = bold
    p.paragraph_format.space_after = Pt(4)
    return p


def add_code(doc, code_text, font_size=10):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = CODE_FONT
    run.font.size = Pt(font_size)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_note(doc, text, is_warning=False):
    p = doc.add_paragraph()
    run = p.add_run("Note: " if not is_warning else "Important: ")
    run.font.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p


def build_document():
    doc = Document()
    doc.add_paragraph()

    # ---- Title ----
    add_heading(doc, "Mindlytics Laravel Deployment Guide", 0)
    add_heading(doc, "Git-based workflow: Local → GitHub → Server", 2)
    add_para(doc, "This guide covers updating the project locally, pushing to GitHub, and deploying changes on the server.")
    doc.add_paragraph()

    # ---- 1. Local (Offline) Updates ----
    add_heading(doc, "1. Local (Offline) Updates", 1)
    add_para(doc, "Open your Laravel project on your computer. After making any changes:")
    add_numbered(doc, "Stage and commit your changes:")
    add_code(doc, "git add .")
    add_code(doc, 'git commit -m "Brief description of the update"')
    add_numbered(doc, "Push updates to GitHub:")
    add_code(doc, "git push origin main")
    add_para(doc, "If your default branch is master instead of main, use: git push origin master")
    add_note(doc, "Never add the .env file to GitHub. It should be listed in .gitignore because it contains sensitive data (database credentials, keys, etc.).", is_warning=True)
    doc.add_paragraph()

    # ---- 2. Server Access & Pull Updates ----
    add_heading(doc, "2. Server Access & Pulling Updates", 1)

    add_heading(doc, "2.1 Connect to the server", 2)
    add_para(doc, "SSH into your server and go to the project directory:")
    add_code(doc, "ssh username@your-server-ip")
    add_code(doc, "cd /path/to/Mindlytics")

    add_heading(doc, "2.2 Verify Git remote", 2)
    add_para(doc, "Check that the project is linked to GitHub:")
    add_code(doc, "git remote -v")
    add_para(doc, "If it is not connected, run:")
    add_code(doc, "git init")
    add_code(doc, "git remote add origin https://github.com/mohamedhany-web/mindlytics.git")

    add_heading(doc, "2.3 Pull latest code from GitHub", 2)
    add_para(doc, "Fetch and reset to match the remote branch exactly:")
    add_code(doc, "git fetch origin")
    add_code(doc, "git reset --hard origin/main")
    add_para(doc, "Using reset --hard ensures all project files match the latest version on GitHub. Your local .env on the server is not tracked by Git, so it will not be overwritten.")
    doc.add_paragraph()

    # ---- 3. .env on server ----
    add_heading(doc, "3. The .env File on the Server", 1)
    add_para(doc, "Do not delete or overwrite the .env file on the server.")
    add_para(doc, "After each update from GitHub, verify that database connection values are correct:")
    add_bullet(doc, "DB_HOST")
    add_bullet(doc, "DB_DATABASE")
    add_bullet(doc, "DB_USERNAME")
    add_bullet(doc, "DB_PASSWORD")
    add_para(doc, "You usually only need to edit .env when you add new configuration options to the project.")
    doc.add_paragraph()

    # ---- 4. Composer ----
    add_heading(doc, "4. Update Composer Dependencies", 1)
    add_para(doc, "If you added or changed PHP packages (composer.json), run on the server:")
    add_code(doc, "composer install --no-dev --optimize-autoloader")
    add_para(doc, "Use --no-dev in production to avoid installing development-only packages.")
    doc.add_paragraph()

    # ---- 5. Migrations ----
    add_heading(doc, "5. Database Migrations", 1)
    add_para(doc, "If there are new or changed migrations (or seeders), run:")
    add_code(doc, "php artisan migrate --force")
    add_para(doc, "The --force flag is required in production to skip confirmation prompts.")
    doc.add_paragraph()

    # ---- 6. Cache ----
    add_heading(doc, "6. Clear and Rebuild Cache", 1)
    add_para(doc, "Laravel caches config, routes, and views. After every deployment run:")
    add_code(doc, "php artisan config:cache")
    add_code(doc, "php artisan route:cache")
    add_code(doc, "php artisan view:cache")
    add_para(doc, "This ensures changes take effect immediately and improves performance.")
    doc.add_paragraph()

    # ---- 7. Summary ----
    add_heading(doc, "7. Workflow Summary", 1)
    add_para(doc, "Quick reference for the full deployment flow.")
    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Step"
    hdr[1].text = "Action"
    for c in hdr:
        c.paragraphs[0].runs[0].font.bold = True
    data = [
        ("Local", "Edit code -> git add . -> git commit -m \"...\" -> git push origin main"),
        ("Server", "git fetch origin && git reset --hard origin/main"),
        ("Packages", "composer install --no-dev --optimize-autoloader (if needed)"),
        ("Database", "php artisan migrate --force (if there are new migrations)"),
        ("Cache", "php artisan config:cache && route:cache && view:cache"),
        (".env", "Keep server .env as-is; never commit it to GitHub."),
    ]
    for i, (step, action) in enumerate(data, 1):
        row = table.rows[i]
        row.cells[0].text = step
        row.cells[1].text = action
    doc.add_paragraph()
    add_note(doc, "Keep .env only on the server and never add it to GitHub.", is_warning=True)

    return doc


def main():
    base = Path(__file__).resolve().parent.parent
    output_path = base / OUTPUT_NAME
    doc = build_document()
    doc.save(str(output_path))
    try:
        print("Created:", str(output_path))
    except UnicodeEncodeError:
        print("Done. File saved:", OUTPUT_NAME)
    return 0


if __name__ == "__main__":
    sys.exit(main())
