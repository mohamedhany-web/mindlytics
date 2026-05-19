#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
دليل شراكة Mindlytics — Word عربي RTL مع دعم المصطلحات الإنجليزية (LTR).

التشغيل:
  pip install -r scripts/requirements-docx.txt
  python scripts/generate_platform_business_docx.py
"""

from __future__ import annotations

import argparse
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ─── خطوط ───
FONT_AR_TITLE = "Traditional Arabic"
FONT_AR_HEADING = "Arabic Typesetting"
FONT_AR_BODY = "Tahoma"
FONT_EN = "Segoe UI"  # للمصطلحات الإنجليزية داخل فقرة عربية

COLOR_PRIMARY = RGBColor(0x1E, 0x3A, 0x8A)
COLOR_ACCENT = RGBColor(0x03, 0x7B, 0xA8)
COLOR_MUTED = RGBColor(0x64, 0x74, 0x8B)
COLOR_PARTNER = RGBColor(0x05, 0x96, 0x69)  # أخضر — قيمة الشريك

# نمط: *عربي* | (English) | نص عادي
TAG_EN = re.compile(r"\(([A-Za-z0-9][^)]*)\)")

# ترتيب عناصر w:pPr حسب OOXML (مهم لـ Word)
_PPR_ORDER = (
    "w:adjustRightInd",
    "w:snapToGrid",
    "w:spacing",
    "w:ind",
    "w:contextualSpacing",
    "w:mirrorIndents",
    "w:suppressOverlap",
    "w:jc",
    "w:textDirection",
    "w:textAlignment",
    "w:textboxTightWrap",
    "w:outlineLvl",
    "w:divId",
    "w:cnfStyle",
    "w:rPr",
    "w:sectPr",
    "w:pPrChange",
)


def _ppr_insert_before(p_pr, element: OxmlElement) -> None:
    """إدراج عنصر pPr قبل العناصر ذات الأولوية الأعلى."""
    for tag in _PPR_ORDER:
        ref = p_pr.find(qn(tag))
        if ref is not None:
            ref.addprevious(element)
            return
    p_pr.append(element)


def _ppr_ensure_child(p_pr, tag: str) -> OxmlElement:
    el = p_pr.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        _ppr_insert_before(p_pr, el)
    return el


def apply_rtl_paragraph(paragraph, *, rtl: bool = True) -> None:
    """
    تفعيل RTL حقيقي في Word:
    w:bidi + w:jc right + محاذاة يمين (بدون w:val على bidi — عنصر منطقي فقط).
    """
    p_pr = paragraph._element.get_or_add_pPr()

    if rtl:
        bidi = p_pr.find(qn("w:bidi"))
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            _ppr_insert_before(p_pr, bidi)
        elif qn("w:val") in bidi.attrib:
            del bidi.attrib[qn("w:val")]

        jc = _ppr_ensure_child(p_pr, "w:jc")
        jc.set(qn("w:val"), "right")

        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf = paragraph.paragraph_format
        pf.right_indent = pf.right_indent  # يحفّز تطبيق RTL في بعض إصدارات Word
    else:
        bidi = p_pr.find(qn("w:bidi"))
        if bidi is not None:
            p_pr.remove(bidi)
        jc = _ppr_ensure_child(p_pr, "w:jc")
        jc.set(qn("w:val"), "left")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _run_font(run, font: str, size: float, *, bold=False, color=None, ltr: bool = False) -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:cs"), font)

    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)

    rtl_el = r_pr.find(qn("w:rtl"))
    if ltr:
        if rtl_el is None:
            rtl_el = OxmlElement("w:rtl")
            r_pr.append(rtl_el)
        rtl_el.set(qn("w:val"), "0")
        lang.set(qn("w:val"), "en-US")
        lang.set(qn("w:eastAsia"), "en-US")
        lang.set(qn("w:bidi"), "ar-SA")
    else:
        if rtl_el is not None:
            r_pr.remove(rtl_el)
        cs = r_pr.find(qn("w:cs"))
        if cs is None:
            cs = OxmlElement("w:cs")
            r_pr.append(cs)
        lang.set(qn("w:val"), "ar-SA")
        lang.set(qn("w:bidi"), "ar-SA")
        lang.set(qn("w:eastAsia"), "ar-SA")

    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_mixed_text(paragraph, text: str, *, size=12, bold=False, color=None) -> None:
    """فقرة RTL: العربي يميناً، ما بين أقواس إنجليزي يُعرض LTR بخط Segoe UI."""
    pos = 0
    for m in TAG_EN.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            _run_font(run, FONT_AR_BODY, size, bold=bold, color=color, ltr=False)
        en = m.group(1)
        run = paragraph.add_run(f" ({en}) ")
        _run_font(run, FONT_EN, size - 0.5, bold=bold, color=color or COLOR_ACCENT, ltr=True)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _run_font(run, FONT_AR_BODY, size, bold=bold, color=color, ltr=False)


def add_para(doc, text: str, *, size=12, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    apply_rtl_paragraph(p, rtl=True)
    add_mixed_text(p, text, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_heading(doc, text: str, level: int = 1):
    cfg = {
        1: (FONT_AR_TITLE, 22, COLOR_PRIMARY, 12),
        2: (FONT_AR_HEADING, 17, COLOR_ACCENT, 9),
        3: (FONT_AR_HEADING, 14, COLOR_PRIMARY, 7),
    }
    font, size, color, sa = cfg.get(level, cfg[3])
    return add_para(doc, text, size=size, bold=True, color=color, space_after=sa)


def add_label(doc, text: str):
    return add_para(doc, text, size=11, bold=True, color=COLOR_PARTNER, space_after=4)


def add_bullets(doc, items: list[str]):
    """قائمة نقطية RTL — بدون نمط List Bullet الإنجليزي الذي يفرض LTR."""
    for item in items:
        p = doc.add_paragraph()
        apply_rtl_paragraph(p, rtl=True)
        p.paragraph_format.right_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after = Pt(3)
        # RLM + نقطة: تظهر على يمين السطر في Word RTL
        bullet_run = p.add_run("\u200f\u2022\t")
        _run_font(bullet_run, FONT_AR_BODY, 11.5, bold=True)
        add_mixed_text(p, item, size=11.5)


def add_module(doc, mod: dict) -> None:
    """قسم واحد: عنوان + وصف + قيمة للشريك + مؤشرات."""
    add_heading(doc, mod["title"], 2)
    if mod.get("subtitle"):
        add_para(doc, mod["subtitle"], size=11, color=COLOR_MUTED, space_after=6)

    add_label(doc, "ماذا يفعل هذا الجزء؟")
    add_para(doc, mod["summary"], space_after=6)

    add_label(doc, "القدرات التفصيلية")
    add_bullets(doc, mod["features"])

    add_label(doc, "لماذا يهم الشريك؟")
    add_bullets(doc, mod["partner_value"])

    if mod.get("metrics"):
        add_label(doc, "مؤشرات يمكن متابعتها")
        add_bullets(doc, mod["metrics"])

    doc.add_paragraph()


# ─── محتوى الأقسام (منظور شريك) ───
PLATFORM_MODULES: list[dict] = [
    {
        "title": "١. الواجهة العامة والموقع التسويقي (Public Website)",
        "subtitle": "أول نقطة لمس يراها العميل والشريك قبل التسجيل.",
        "summary": (
            "الواجهة العامة هي واجهة العلامة التجارية: تعرض رسالة الأكاديمية، الكورسات، المسارات، "
            "المدربين، والإحصائيات (متعلمون، شهادات، دقائق مشاهدة). تدعم العربية والإنجليزية "
            "(Arabic / English) لتوسيع قاعدة العملاء."
        ),
        "features": [
            "صفحة رئيسية (Landing) بأقسام: بطل، إحصائيات، كورسات مميزة، مسارات، بورتفوليو، ودعوة للتسجيل.",
            "كتالوج كورسات ومسارات تعليمية (Learning Paths) قابلة للتصفح دون دخول.",
            "صفحة المدربين وبروفايلاتهم.",
            "معرض أعمال عام (Portfolio) للشركات والزوار.",
            "تسجيل دخول وإنشاء حساب طالب.",
            "صفحة «من نحن» قابلة للإدارة من لوحة الإدارة.",
            "مدونة (Blog) ورسائل تواصل (Contact).",
            "باقات تسعير (Packages) معروضة للجمهور.",
        ],
        "partner_value": [
            "قناة اكتساب عملاء منخفضة التكلفة مقارنة بالإعلان فقط.",
            "مصداقية فورية عبر الأرقام الحية (طلاب، شهادات، كورسات).",
            "قصة علامة موحّدة من أول زيارة حتى الدفع.",
            "جاهزية للتوسع الإقليمي بلغتين.",
        ],
        "metrics": ["زوار فريدون", "معدل التحويل للتسجيل", "طلبات التواصل غير المقروءة"],
    },
    {
        "title": "٢. عملية الشراء والدفع (Checkout & Payments)",
        "summary": (
            "مسار مالي واضح من اختيار الكورس أو المسار حتى تأكيد الطلب، مع دعم بوابات دفع "
            "وتقسيط يزيد معدل الإغلاق."
        ),
        "features": [
            "صفحة إتمام طلب (Checkout) لكل كورس أو مسار.",
            "تتبع الطلبات (Orders) من لوحة الطالب والإدارة.",
            "فواتير (Invoices) ومدفوعات (Payments) ومعاملات (Transactions).",
            "محفظة طالب (Wallet) لرصيد وحركات مالية.",
            "بوابات دفع متعددة من مركز المحاسبة.",
            "خطط تقسيط (Installment Plans) وأتفاقيات سداد.",
            "كوبونات وخصومات وبرامج إحالة.",
        ],
        "partner_value": [
            "إيرادات قابلة للتنبؤ وتقارير مالية جاهزة.",
            "تقسيط يفتح شريحة عملاء أوسع دون خصم على الجودة.",
            "شفافية كاملة للشريك في كل عملية بيع.",
        ],
        "metrics": ["إجمالي المبيعات", "طلبات معلّقة", "نسبة التقسيط", "متوسط قيمة الطلب"],
    },
    {
        "title": "٣. المسارات التعليمية (Learning Paths)",
        "summary": (
            "المسار يجمع عدة كورسات ومهارات في رحلة واحدة — منتج أعلى قيمة من كورس منفرد."
        ),
        "features": [
            "تعريف مسارات من الإدارة وربط الكورسات بها.",
            "تعيين مدربين لمسارات محددة.",
            "تسجيل طلاب في المسار ومتابعة حالة الطلب (موافقة / نشط).",
            "لوحة طالب للمسار النشط وتقدم المراحل.",
            "تقييمات المسارات (Learning Path Reviews).",
        ],
        "partner_value": [
            "رفع متوسط قيمة البيع (ARPU) عبر باقات مسارات.",
            "تقليل التسرب: الطالب ملتزم بخطة طويلة.",
            "تمييز تجاري عن منصات «كورس واحد فقط».",
        ],
        "metrics": ["مسجّلون نشطون في المسارات", "معدل إكمال المسار", "تقييمات المسارات"],
    },
    {
        "title": "٤. الكورسات المسجّلة أونلاين (Recorded Online Courses)",
        "summary": (
            "تعلّم ذاتي عبر فيديوهات ومنهج منظم — قابل للتوسع بدون سقف حضور físico."
        ),
        "features": [
            "كتالوج كورسات متقدمة (Advanced Courses) مع دروس ومحاضرات.",
            "مشغل تعلّم موحّد مع تتبع التقدم (Progress).",
            "موارد وملفات لكل محاضرة.",
            "مجموعات دراسية (Groups) داخل الكورس.",
            "واجبات (Assignments) ومشاريع.",
            "امتحانات (Exams) وبنوك أسئلة (Question Banks).",
            "تسجيلات أونلاين وإدارة التسجيل من الإدارة.",
        ],
        "partner_value": [
            "هامش ربح أعلى بعد إنتاج المحتوى مرة واحدة.",
            "وصول لطلاب خارج المدن دون فروع إضافية.",
            "قابلية التوسع السريع في المحتوى.",
        ],
        "metrics": ["دقائق مشاهدة", "نسبة إكمال الكورس", "نتائج الامتحانات"],
    },
    {
        "title": "٥. الكورسات الحضورية — أوفلاين (Offline Courses)",
        "summary": (
            "برامج ميدانية بقاعات وفروع: حضور، مجموعات، حجوزات، واتفاقيات مدربين."
        ),
        "features": [
            "إدارة مواقع التدريب (Locations).",
            "كورسات أوفلاين بمجموعات وجداول.",
            "تسجيل طلاب وحجوزات أوفلاين (مع موافقة الإدارة).",
            "حضور محاضرات وجلسات (Attendance).",
            "موارد وتسجيلات محاضرات داخل المنصة.",
            "تقارير طلاب لكل كورس حضوري.",
            "اتفاقيات مالية مع مدربي الأوفلاين.",
            "أنشطة وتكاملات إضافية للبرنامج الحضوري.",
        ],
        "partner_value": [
            "منتج premium بأسعار أعلى وولاء أقوى.",
            "تحكم كامل في السعة والفروع.",
            "ربط الإيراد بالحضور الفعلي — أقل مخاطرة «بيع وهمي».",
        ],
        "metrics": ["مقاعد محجوزة / مشغولة", "نسبة الحضور", "حجوزات معلّقة"],
    },
    {
        "title": "٦. الكورسات الأونلاين الجماعية (Live Online Group Courses)",
        "summary": (
            "جلسات مباشرة أو مجموعات أونلاين ببوابة طالب مفعّلة — بين المرونة والتفاعل."
        ),
        "features": [
            "إنشاء كورس أونلاين جماعي من الإدارة.",
            "تسجيل طالب بالبريد وتفعيل بوابة الطالب.",
            "حجوزات أونلاين (Online Bookings) بقناة منفصلة عن الأوفلاين.",
            "لوحة مدرب لمجموعات الأونلاين وطلاب نشطين.",
            "تقويم جلسات مشترك.",
        ],
        "partner_value": [
            "توسيع الجغرافيا دون فتح فرع.",
            "هامش وسط بين المسجّل والحضوري.",
            "مرونة في الجداول والمواسم.",
        ],
        "metrics": ["طلاب نشطون أونلاين جماعي", "حجوزات أونلاين معلّقة"],
    },
    {
        "title": "٧. لوحة الطالب (Student Portal)",
        "summary": (
            "مركز تجربة المتعلّم: من الاكتشاف حتى الشهادة والبورتفوليو."
        ),
        "features": [
            "لوحة تحكم: تقدم، كورسات نشطة، إحصائيات سريعة.",
            "تصفح وشراء كورسات جديدة.",
            "كورساتي: تعلّم، واجبات، امتحانات.",
            "مجموعاتي (Groups).",
            "مسار تعليمي نشط.",
            "فصل واضح: أوفلاين / أونلاين جماعي.",
            "طلباتي وفواتيري ومحفظتي.",
            "شهاداتي وإنجازاتي وشاراتي.",
            "تقويم وإشعارات.",
            "مجتمع الذكاء الاصطناعي (مسابقات، Datasets، Model Zoo).",
            "الملف الشخصي والإعدادات.",
        ],
        "partner_value": [
            "احتفاظ عالٍ بالعميل داخل منظومة واحدة.",
            "تقليل شكاوى الدعم — كل شيء في مكان واحد.",
            "بيانات سلوك تعلّم تغذي قرارات المنتج.",
        ],
        "metrics": ["نشاط أسبوعي", "نسبة إكمال", "رضا من التقييمات"],
    },
    {
        "title": "٨. لوحة المدرب (Instructor Portal)",
        "summary": (
            "تمكين المدرب من التدريس والمتابعة مع ربط مالي عبر الاتفاقيات والرواتب."
        ),
        "features": [
            "كورساتي المسجّلة وأوفلاين وأونلاين جماعي.",
            "بناء منهج: محاضرات، فيديو، موارد.",
            "مجموعات، واجبات، امتحانات، بنك أسئلة.",
            "حضور وجلسات وتقارير طلاب.",
            "تقويم ومهام من الإدارة.",
            "طلبات للإدارة (انضمام، تعديلات).",
            "بروفايل وعلامة شخصية (Personal Branding).",
        ],
        "partner_value": [
            "جذب مدربين بجودة دون أدوات خارجية.",
            "جودة تعليم مرتبطة بعقود واضحة.",
            "توسيع الكتالوج بسرعة عبر شبكة مدربين.",
        ],
        "metrics": ["عدد طلاب لكل مدرب", "تسليم الواجبات", "تقييمات المدرب"],
    },
    {
        "title": "٩. معرض الأعمال — البورتفوليو (Portfolio)",
        "summary": (
            "جسر بين التعليم والتوظيف: الطالب يعرض مشاريعه والشركة تكتشف المواهب."
        ),
        "features": [
            "رفع مشاريع بعد الكورسات.",
            "معرض عام للزوار والشركات.",
            "إدارة المحتوى من الإدارة (موافقة، ترتيب).",
            "ربط الشهادات والإنجازات بالملف.",
            "مسار: تعلّم → مشروع → شهادة → فرصة عمل.",
        ],
        "partner_value": [
            "قيمة B2B: شركات تدفع أو تتعاون للوصول للمواهب.",
            "تمييز قوي في السوق العربي.",
            "نتائج ملموسة تُسوَّق للطلاب والشركاء.",
        ],
        "metrics": ["مشاريع منشورة", "زيارات المعرض", "تواصل شركات"],
    },
    {
        "title": "١٠. مجتمع البيانات والذكاء الاصطناعي (AI & Data Community)",
        "summary": (
            "طبقة مجتمعية متخصصة: مسابقات، بيانات، نماذج، ومناقشات — تعزز الولاء والعلامة."
        ),
        "features": [
            "مسابقات (Competitions) بإدارة كاملة.",
            "مجموعات بيانات (Datasets) وتقديمات للمراجعة.",
            "Model Zoo لتقديم النماذج.",
            "مناقشات (Discussions) وإشعارات مجتمع.",
            "مساهمون (Contributors) بصلاحيات خاصة.",
            "لوحة مراقبة للإدارة العليا (Super Admin).",
        ],
        "partner_value": [
            "شراكات مع شركات تقنية ومؤسسات بحثية.",
            "محتوى مجاني يجذب قاعدة مستخدمين.",
            "تموضع كمركز عربي للـ AI/Data وليس أكاديمية عامة فقط.",
        ],
        "metrics": ["مشاركون في المسابقات", "تقديمات datasets/models", "نشاط المناقشات"],
    },
    {
        "title": "١١. إدارة المبيعات (Sales Module)",
        "summary": (
            "نظام CRM داخلي: عملاء محتملون، أنشطة، أهداف، عمولات، وتقارير."
        ),
        "features": [
            "تقارير مبيعات وتصدير (Export).",
            "عملاء محتملون (Leads) وسجل أنشطة.",
            "مراقبة KPIs وأهداف الفريق (Targets).",
            "رؤى أداء الموظفين (Insights).",
            "عمولات مبيعات (Commissions).",
            "لوحة موظف مبيعات: leads، CSAT، KPI شخصي.",
        ],
        "partner_value": [
            "نمو إيرادات قابل للقياس والمحاسبة.",
            "حوافز فريق مبيعات شفافة.",
            "رؤية pipeline للشريك قبل إغلاق الصفقات.",
        ],
        "metrics": ["Leads محوّلة", "KPI الشهري", "CSAT", "العمولات المستحقة"],
    },
    {
        "title": "١٢. المحاسبة والمالية (Accounting & Finance)",
        "summary": (
            "مركز مالي متكامل: من الفاتورة حتى شجرة الحسابات ومؤشرات لحظية."
        ),
        "features": [
            "مركز محاسبة (Accounting Hub) ومؤشرات لحظية (Real-time Insights).",
            "شجرة حسابات (Chart of Accounts).",
            "فواتير، مدفوعات، معاملات، محافظ.",
            "مصروفات واشتراكات.",
            "تقسيط: خطط وأتفاقيات سداد ولوحة تقسيط.",
            "رواتب وحسابات مدربين.",
            "اتفاقيات ورواتب موظفين.",
            "بوابات دفع وعمليات البوابة.",
            "تقارير محاسبية شاملة.",
        ],
        "partner_value": [
            "جاهزية Due Diligence للمستثمر أو الشريك المالي.",
            "قرارات تسعير وتوسع مبنية على أرقام.",
            "تقليل الاعتماد على Excel خارجي.",
        ],
        "metrics": ["إيرادات / مصروفات", "ذمم مدينة", "تدفق نقدي", "أرباح حسب المنتج"],
    },
    {
        "title": "١٣. التسويق والولاء (Marketing & Loyalty)",
        "summary": (
            "أدوات نمو: إعلانات، كوبونات، إحالة، ولاء، وهوية المدرب."
        ),
        "features": [
            "إعلانات منبثقة (Pop-up Ads).",
            "علامة شخصية للمدرب (Personal Branding).",
            "كوبونات وخصومات.",
            "برامج إحالة (Referral Programs) ومتابعة الإحالات.",
            "برامج ولاء (Loyalty).",
            "خطط تسويق المشرفين عبر منصات التواصل.",
            "دورات تصميم (Design Cycles) وتسليم مخرجات.",
        ],
        "partner_value": [
            "خفض تكلفة اكتساب العميل (CAC).",
            "حملات قابلة للقياس داخل المنصة.",
            "تسويق منظم للفريق الداخلي والشركاء.",
        ],
        "metrics": ["استخدام الكوبونات", "إحالات ناجحة", "نقاط الولاء"],
    },
    {
        "title": "١٤. الموارد البشرية والموظفون (HR & Employees)",
        "summary": (
            "إدارة الفريق الداخلي: وظائف، مهام، إجازات، رواتب، وخصومات."
        ),
        "features": [
            "سجل موظفين ووظائف (Jobs).",
            "مهام موظفين مع تسليم مخرجات (Deliverables).",
            "طلبات إجازة (Leaves) بموافقة الإدارة.",
            "خصومات رواتب (Deductions).",
            "اتفاقيات موظفين ومدفوعات.",
            "مهام مدربين وطلبات انضمام كمدرب.",
            "إشعارات موظفين منفصلة عن الطلاب.",
            "تقويم وتقارير موظف.",
        ],
        "partner_value": [
            "تشغيل أكاديمية متوسعة دون أنظمة HR منفصلة.",
            "مساءلة الفريق عبر مهام ومؤشرات.",
            "تكلفة تشغيل واضحة للشريك.",
        ],
        "metrics": ["مهام متأخرة", "إجازات معلّقة", "تكلفة رواتب شهرية"],
    },
    {
        "title": "١٥. الرقابة والجودة (Quality Control)",
        "summary": (
            "متابعة أداء الطلاب والمدربين والموظفين وضمان معايير التشغيل."
        ),
        "features": [
            "لوحة رقابة مركزية.",
            "رقابة طلاب: تقدم، نشاط، مخاطر تسرب.",
            "رقابة مدربين: التزام، جودة تسليم.",
            "رقابة موظفين: إنتاجية.",
            "متابعة عمليات يومية (Operations Follow-up).",
        ],
        "partner_value": [
            "حماية سمعة العلامة أمام الشريك والعميل.",
            "تدخل مبكر قبل شكاوى أو استرداد.",
            "معيار جودة قابل للتدقيق في عقود B2B.",
        ],
        "metrics": ["طلاب معرّضون للتسرب", "تأخر تسليمات", "شكاوى مفتوحة"],
    },
    {
        "title": "١٦. الشهادات والتحفيز (Certificates & Gamification)",
        "summary": (
            "إثبات الإنجاز: شهادات رسمية، شارات، إنجازات، وتقييمات."
        ),
        "features": [
            "إصدار ومراجعة شهادات (قائمة، إنشاء، معلّقة).",
            "إنجازات (Achievements) وشارات (Badges).",
            "تقييمات كورسات ومسارات.",
            "عرض الشهادات في بروفايل الطالب والبورتفوليو.",
        ],
        "partner_value": [
            "منتج تسويقي: «شهادة معتمدة» ترفع معدل التحويل.",
            "ثقة الشركات في مهارات الخريجين.",
            "محتوى قابل للمشاركة على LinkedIn ووسائل التواصل.",
        ],
        "metrics": ["شهادات صادرة", "شهادات معلّقة", "متوسط التقييم"],
    },
    {
        "title": "١٧. الورش والاجتماعات (Workshops & Meetings)",
        "summary": (
            "فعاليات تكميلية: ورش عمل وجلسات جماعية تدعم الكورسات والمسارات."
        ),
        "features": [
            "جدولة وإدارة ورش من لوحة الإدارة.",
            "ربط الورش بالتقويم والإشعارات.",
            "تسجيل حضور أو مشاركة حسب الإعداد.",
        ],
        "partner_value": [
            "إيرادات إضافية من فعاليات قصيرة.",
            "تجربة مجتمع حية تعزز الولاء.",
            "شراكات رعاية فعاليات مع شركات.",
        ],
        "metrics": ["ورش منفذة", "حضور الورش", "إيراد الورش"],
    },
    {
        "title": "١٨. إدارة النظام والصلاحيات (System & Permissions)",
        "summary": (
            "حوكمة الوصول: أدوار، صلاحيات، سجلات أمان، وإعدادات المنصة."
        ),
        "features": [
            "إدارة مستخدمين (طلاب، مدربين، موظفين، إداريين).",
            "أدوار (Roles) وصلاحيات (Permissions) وصلاحيات فردية.",
            "إعدادات النظام والهوية البصرية (شعار، ألوان).",
            "سجل نشاط (Activity Log) وسجلات تحقق بخطوتين (2FA Logs).",
            "إحصائيات وأداء المنصة (Performance).",
            "مصادر فيديو (Video Providers).",
        ],
        "partner_value": [
            "أمان وامتثال عند دخول شريك استثماري أو مؤسسي.",
            "مرونة فتح فروع أو فرق بصلاحيات محددة.",
            "تقليل مخاطر تسريب بيانات أو تلاعب.",
        ],
        "metrics": ["محاولات دخول فاشلة", "نشاط إداري حرج", "وقت استجابة المنصة"],
    },
    {
        "title": "١٩. التقارير الشاملة (Reports & Analytics)",
        "summary": (
            "تقارير جاهزة للتصدير: مستخدمون، كورسات، مالية، أكاديمية، وأنشطة."
        ),
        "features": [
            "لوحة تقارير (Reports Dashboard).",
            "تقارير مستخدمين وكورسات.",
            "تقارير مالية وأكاديمية.",
            "تقارير أنشطة وتقرير شامل (Comprehensive).",
            "تصدير Excel حيث متاح.",
        ],
        "partner_value": [
            "اجتماعات مجلس إدارة وشركاء بدون تجهيز يدوي.",
            "شفافية أداء الأكاديمية في عقد شراكة.",
            "أساس لاتفاقيات أداء (SLA) مع فرق داخلية.",
        ],
        "metrics": ["نمو المستخدمين", "إيراد شهري", "معدل إكمال أكاديمي"],
    },
    {
        "title": "٢٠. الاتفاقيات والسحوبات (Agreements & Withdrawals)",
        "summary": (
            "عقود مالية مع المدربين والموظفين وطلبات سحب الأرصدة."
        ),
        "features": [
            "اتفاقيات مدربين (Instructor Agreements).",
            "اتفاقيات موظفين ودفعات مرتبطة.",
            "طلبات سحب (Withdrawal Requests) بموافقة الإدارة.",
            "ربط الاتفاقيات بالمحاسبة والرواتب.",
        ],
        "partner_value": [
            "علاقات قانونية ومالية منظمة مع الموردين (مدربين).",
            "ثقة المدرب في سياسة دفع واضحة.",
            "تدفق نقدي خاضع للرقابة.",
        ],
        "metrics": ["سحوبات معلّقة", "التزامات اتفاقيات", "مدفوعات شهرية للمدربين"],
    },
    {
        "title": "٢١. الرسائل والتواصل الداخلي (Messages & Notifications)",
        "summary": (
            "قنوات تواصل: إشعارات للطلاب والموظفين، رسائل إدارية، وبريد المنصة."
        ),
        "features": [
            "إشعارات طلاب (Navbar / Email).",
            "إشعارات موظفين منفصلة.",
            "رسائل إدارية (Messages).",
            "إشعارات مجتمع الذكاء الاصطناعي.",
            "تقويم موحّد مع تذكيرات.",
        ],
        "partner_value": [
            "تقليل فقدان العميل بسبب ضعف التواصل.",
            "حملات تنشيط (Re-engagement) داخل المنصة.",
            "تجربة احترافية تليق بعلامة متوسعة.",
        ],
        "metrics": ["معدل فتح الإشعارات", "رسائل غير مقروءة"],
    },
]

USER_ROWS = [
    ("الطالب (Student)", "متعلّم", "شراء، تعلّم، شهادة، بورتفوليو، مجتمع AI"),
    ("المدرب (Instructor / Teacher)", "مقدّم محتوى", "كورسات، حضور، تقييم، اتفاقيات مالية"),
    ("الإداري (Admin)", "تشغيل يومي", "محتوى، مبيعات، محاسبة، تسجيلات، جودة"),
    ("المدير الأعلى (Super Admin)", "حوكمة", "مجتمع AI، صلاحيات، إعدادات حساسة"),
    ("موظف المبيعات (Sales)", "إيرادات", "Leads، KPI، عمولات، CSAT"),
    ("مشرف التسويق (Moderator)", "حضور رقمي", "خطط منصات، تصميم، تقويم محتوى"),
    ("الموظف العام (Employee)", "تنفيذ", "مهام، إجازات، راتب، اتفاقيات"),
    ("الزائر (Visitor)", "اكتساب", "تصفح، تسجيل، دفع، معرض أعمال"),
    ("الشركة (Company)", "توظيف", "بورتفوليو، اكتشاف مواهب، شراكات B2B"),
]


def _table_rtl(table) -> None:
    """عكس أعمدة الجدول بصرياً (العمود الأول يميناً)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    if tbl_pr.find(qn("w:bidiVisual")) is None:
        tbl_pr.append(OxmlElement("w:bidiVisual"))


def add_users_table(doc):
    headers = ["الدور", "الوظيفة", "قيمة للشريك"]
    table = doc.add_table(rows=1 + len(USER_ROWS), cols=3)
    table.style = "Table Grid"
    _table_rtl(table)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        apply_rtl_paragraph(p, rtl=True)
        add_mixed_text(p, h, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        sh = OxmlElement("w:shd")
        sh.set(qn("w:fill"), "1E3A8A")
        cell._tc.get_or_add_tcPr().append(sh)
    for ri, row in enumerate(USER_ROWS, 1):
        for ci, txt in enumerate(row):
            p = table.rows[ri].cells[ci].paragraphs[0]
            apply_rtl_paragraph(p, rtl=True)
            add_mixed_text(p, txt, size=10.5)
    doc.add_paragraph()


def _configure_document_rtl(doc: Document) -> None:
    """RTL على مستوى المستند: docDefaults + Normal + sectPr + لغة عربية."""
    styles_el = doc.styles.element

    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_el.insert(0, doc_defaults)

    p_pr_default = doc_defaults.find(qn("w:pPrDefault"))
    if p_pr_default is None:
        p_pr_default = OxmlElement("w:pPrDefault")
        doc_defaults.append(p_pr_default)

    default_p_pr = p_pr_default.find(qn("w:pPr"))
    if default_p_pr is None:
        default_p_pr = OxmlElement("w:pPr")
        p_pr_default.append(default_p_pr)

    if default_p_pr.find(qn("w:bidi")) is None:
        default_p_pr.append(OxmlElement("w:bidi"))
    default_jc = default_p_pr.find(qn("w:jc"))
    if default_jc is None:
        default_jc = OxmlElement("w:jc")
        default_p_pr.append(default_jc)
    default_jc.set(qn("w:val"), "right")

    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    if r_pr_default is None:
        r_pr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(r_pr_default)
    default_r_pr = r_pr_default.find(qn("w:rPr"))
    if default_r_pr is None:
        default_r_pr = OxmlElement("w:rPr")
        r_pr_default.append(default_r_pr)
    drfonts = default_r_pr.find(qn("w:rFonts"))
    if drfonts is None:
        drfonts = OxmlElement("w:rFonts")
        default_r_pr.append(drfonts)
    drfonts.set(qn("w:ascii"), FONT_AR_BODY)
    drfonts.set(qn("w:hAnsi"), FONT_AR_BODY)
    drfonts.set(qn("w:cs"), FONT_AR_BODY)
    dlang = default_r_pr.find(qn("w:lang"))
    if dlang is None:
        dlang = OxmlElement("w:lang")
        default_r_pr.append(dlang)
    dlang.set(qn("w:val"), "ar-SA")
    dlang.set(qn("w:bidi"), "ar-SA")
    dlang.set(qn("w:eastAsia"), "ar-SA")

    normal = doc.styles["Normal"]
    normal.font.name = FONT_AR_BODY
    normal.font.size = Pt(12)
    normal_p_pr = normal._element.get_or_add_pPr()
    if normal_p_pr.find(qn("w:bidi")) is None:
        normal_p_pr.append(OxmlElement("w:bidi"))
    njc = normal_p_pr.find(qn("w:jc"))
    if njc is None:
        njc = OxmlElement("w:jc")
        normal_p_pr.append(njc)
    njc.set(qn("w:val"), "right")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        sect_pr = OxmlElement("w:sectPr")
        body.append(sect_pr)
    if sect_pr.find(qn("w:bidi")) is None:
        sect_pr.append(OxmlElement("w:bidi"))
    if sect_pr.find(qn("w:rtlGutter")) is None:
        sect_pr.append(OxmlElement("w:rtlGutter"))

    for section in doc.sections:
        ssp = section._sectPr
        if ssp.find(qn("w:bidi")) is None:
            ssp.append(OxmlElement("w:bidi"))

    settings = doc.settings.element
    theme_lang = settings.find(qn("w:themeFontLang"))
    if theme_lang is None:
        theme_lang = OxmlElement("w:themeFontLang")
        settings.append(theme_lang)
    theme_lang.set(qn("w:val"), "en-US")
    theme_lang.set(qn("w:eastAsia"), "en-US")
    theme_lang.set(qn("w:bidi"), "ar-SA")


def _rtlize_all_paragraphs(doc: Document) -> None:
    """تطبيق RTL على كل فقرة (بما فيها الفارغة بعد page break)."""
    for p in doc.paragraphs:
        apply_rtl_paragraph(p, rtl=True)
    for table in doc.tables:
        _table_rtl(table)
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    apply_rtl_paragraph(p, rtl=True)


def build_document_com(output_path: Path) -> Path:
    """إنشاء المستند عبر Microsoft Word — RTL صحيح للترقيم والنقاط."""
    from word_rtl_com import WordRtlDocument

    today = datetime.now().strftime("%Y/%m/%d")
    w = WordRtlDocument()

    try:
        w.paragraph("Mindlytics", size=34, bold=True, color=0x8A3A1E, font_bi="Segoe UI", space_after=4)
        w.paragraph("أكاديمية البرمجة والذكاء الاصطناعي", size=20, bold=True, color=0xA87B03, space_after=10)
        w.paragraph("دليل الشراكة — عرض المنصة الكامل", size=16, bold=True, space_after=8)
        w.paragraph(
            "وثيقة موجّهة للشركاء الاستراتيجيين والمستثمرين وصنّاع القرار. "
            "عربي من اليمين إلى اليسار — المصطلحات الإنجليزية بين أقواس.",
            size=11,
            color=0x8B7464,
            space_after=20,
        )
        w.paragraph(f"تاريخ الإصدار: {today}", size=10, color=0x8B7464)
        w.page_break()

        w.heading("مقدمة للشريك", 1)
        w.paragraph(
            "منصة Mindlytics ليست موقع فيديوهات فقط، بل نظاماً بيئياً متكاملاً: تعليم عربي "
            "في البرمجة والذكاء الاصطناعي، إدارة مؤسسية، مبيعات، محاسبة، مجتمع تقني، وربط "
            "بسوق العمل عبر البورتفوليو. كل قسم يوضّح: ماذا يفعل؟ ولماذا يهم الشريك؟",
        )
        w.label("لماذا هذا الدليل؟")
        w.bullets(
            [
                "فهم كل جزء من المنصة قبل الشراكة أو الاستثمار.",
                "ربط القدرات بمصادر الإيراد وتقليل المخاطر التشغيلية.",
                "لغة موحّدة بين الفريق العربي والمصطلحات العالمية (KPI، B2B، Leads).",
            ]
        )
        w.page_break()

        w.heading("خريطة المنظومة — نظرة الشريك", 1)
        w.paragraph(
            "المنصة تتكون من أربع طبقات: (1) اكتساب العميل، (2) التعليم والتجربة، "
            "(3) النمو والإيراد، (4) التميز والولاء."
        )
        w.bullets(
            [
                "طبقة العميل: واجهة عامة، شراء، تسجيل.",
                "طبقة التعليم: كورسات مسجّلة، أوفلاين، أونلاين جماعي، مسارات.",
                "طبقة التشغيل: إدارة، موارد بشرية، جودة، تقارير.",
                "طبقة القيمة المضافة: بورتفوليو، مجتمع الذكاء الاصطناعي، شهادات.",
            ]
        )
        w.page_break()

        w.heading("أدوار المستخدمين وقيمتها للشريك", 1)
        w.paragraph("كل دور يرى لوحة وصلاحيات مختلفة:")
        w.table(
            ["الدور", "الوظيفة", "قيمة للشريك"],
            USER_ROWS,
        )
        w.page_break()

        w.heading("شرح تفصيلي لكل جزء في المنصة", 1)
        w.paragraph(
            "كل قسم: الوصف، القدرات، قيمة الشريك، ومؤشرات مقترحة.",
            space_after=10,
        )
        for mod in PLATFORM_MODULES:
            w.module(mod)

        w.heading("نموذج الشراكة والإيرادات — ملخص للشريك", 1)
        w.bullets(
            [
                "إيرادات مباشرة: كورسات، مسارات، أوفلاين، ورش.",
                "إيرادات متكررة: اشتراكات وتقسيط.",
                "نمو عضوي: إحالة وولاء وكوبونات.",
                "إيرادات B2B: شركات عبر البورتفوليو ورعاية مجتمع الذكاء الاصطناعي.",
                "كفاءة تشغيل: منصة واحدة — تكلفة أقل للشريك.",
            ]
        )

        w.heading("الخلاصة", 1)
        w.paragraph(
            "الشراكة مع Mindlytics تعني منظومة جاهزة للتوسع: منتج تعليمي قوي، عمليات مؤسسية، "
            "وتمييز في السوق العربي عبر البورتفوليو ومجتمع الذكاء الاصطناعي.",
        )
        w.paragraph("— نهاية دليل الشراكة —", size=11, color=0x8B7464, space_after=0)

        w.save(output_path)
    finally:
        w.close()

    return output_path.resolve()


def build_document_docx(output_path: Path) -> Path:
    """نسخة احتياطية عبر python-docx (قد لا يعرض RTL بشكل مثالي في Word)."""
    doc = Document()
    _configure_document_rtl(doc)
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)

    today = datetime.now().strftime("%Y/%m/%d")

    # غلاف
    p = doc.add_paragraph()
    apply_rtl_paragraph(p, rtl=True)
    r = p.add_run("Mindlytics")
    _run_font(r, FONT_EN, 34, bold=True, color=COLOR_PRIMARY, ltr=True)
    p.paragraph_format.space_after = Pt(4)

    add_para(doc, "أكاديمية البرمجة والذكاء الاصطناعي", size=20, bold=True, color=COLOR_ACCENT, space_after=10)
    add_para(doc, "دليل الشراكة — عرض المنصة الكامل", size=16, bold=True, space_after=8)
    add_para(
        doc,
        "وثيقة موجّهة للشركاء الاستراتيجيين والمستثمرين وصنّاع القرار\n"
        "عربي من اليمين إلى اليسار (RTL) — المصطلحات الإنجليزية بين أقواس وتُعرض باتجاهها الصحيح (LTR)",
        size=11,
        color=COLOR_MUTED,
        space_after=20,
    )
    add_para(doc, f"تاريخ الإصدار: {today}", size=10, color=COLOR_MUTED)
    doc.add_page_break()

    # مقدمة للشريك
    add_heading(doc, "مقدمة للشريك", 1)
    add_para(
        doc,
        "منصة (Mindlytics) ليست موقع فيديوهات فقط، بل نظاماً بيئياً متكاملاً: تعليم عربي "
        "في البرمجة والذكاء الاصطناعي، إدارة مؤسسية، مبيعات، محاسبة، مجتمع تقني، وربط "
        "بسوق العمل عبر (Portfolio). كل قسم في هذا الدليل يوضّح: ماذا يفعل؟ ولماذا يهم "
        "الشريك؟ — دون دخول في تفاصيل تقنية التطوير.",
    )
    add_label(doc, "لماذا هذا الدليل؟")
    add_bullets(
        doc,
        [
            "فهم كل جزء من المنصة قبل الشراكة أو الاستثمار.",
            "ربط القدرات بمصادر الإيراد وتقليل المخاطر التشغيلية.",
            "لغة موحّدة بين الفريق العربي والمصطلحات العالمية (KPI، B2B، Leads…).",
        ],
    )
    doc.add_page_break()

    # خريطة المنظومة
    add_heading(doc, "خريطة المنظومة — نظرة الشريك", 1)
    add_para(
        doc,
        "المنصة تتكون من أربع طبقات مترابطة: (1) اكتساب العميل عبر الموقع العام، "
        "(2) التعليم والتجربة (طالب + مدرب)، (3) النمو والإيراد (مبيعات + تسويق + محاسبة)، "
        "(4) التميز والولاء (شهادات + بورتفوليو + مجتمع AI + جودة).",
    )
    add_bullets(
        doc,
        [
            "طبقة العميل: واجهة عامة، شراء، تسجيل.",
            "طبقة التعليم: كورسات مسجّلة، أوفلاين، أونلاين جماعي، مسارات.",
            "طبقة التشغيل: إدارة، HR، جودة، تقارير.",
            "طبقة القيمة المضافة: بورتفوليو، مجتمع AI، شهادات.",
        ],
    )
    doc.add_page_break()

    # جدول المستخدمين
    add_heading(doc, "أدوار المستخدمين وقيمتها للشريك", 1)
    add_para(doc, "كل دور يرى لوحة وصلاحيات مختلفة — الجدول يلخّص العلاقة بالأعمال:")
    add_users_table(doc)
    doc.add_page_break()

    # كل الأقسام
    add_heading(doc, "شرح تفصيلي لكل جزء في المنصة", 1)
    add_para(
        doc,
        "يأتي كل قسم بالترتيب التالي: الوصف، القدرات، قيمة الشريك، ومؤشرات مقترحة للمتابعة. "
        "المصطلحات الإنجليزية بين أقواس منسّقة تلقائياً للقراءة الصحيحة.",
        space_after=10,
    )

    for mod in PLATFORM_MODULES:
        add_module(doc, mod)

    # نموذج شراكة
    add_heading(doc, "نموذج الشراكة والإيرادات — ملخص للشريك", 1)
    add_bullets(
        doc,
        [
            "إيرادات مباشرة: كورسات، مسارات، أوفلاين، ورش.",
            "إيرادات متكررة: اشتراكات وتقسيط.",
            "نمو عضوي: إحالة وولاء وكوبونات.",
            "إيرادات B2B: شركات عبر (Portfolio) ورعاية مجتمع (AI).",
            "كفاءة تشغيل: منصة واحدة بدل أنظمة متفرقة — تكلفة أقل للشريك.",
        ],
    )

    add_heading(doc, "الخلاصة", 1)
    add_para(
        doc,
        "الاستثمار أو الشراكة مع (Mindlytics) يعني الدخول في منظومة جاهزة للتوسع: "
        "منتج تعليمي قوي، عمليات مؤسسية، وتمييز في السوق العربي عبر البورتفوليو ومجتمع "
        "الذكاء الاصطناعي. هذا الدليل يغطي كل جزء تشغيلي في المنصة من منظور يخدم "
        "قرار الشريك — وليس الفريق التقني.",
    )
    add_para(doc, "— نهاية دليل الشراكة —", size=11, color=COLOR_MUTED, space_after=0)

    _rtlize_all_paragraphs(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def main():
    import sys

    parser = argparse.ArgumentParser(description="Mindlytics partner guide DOCX (RTL via Microsoft Word)")
    root = Path(__file__).resolve().parents[1]
    default_out = root / "storage" / "docs" / "Mindlytics-Partner-Guide-Ar.docx"
    parser.add_argument("-o", "--output", type=Path, default=default_out)
    parser.add_argument(
        "--fallback-docx",
        action="store_true",
        help="استخدام python-docx بدل Word (غير موصى به للعربية)",
    )
    args = parser.parse_args()
    out_path = args.output.resolve()

    scripts_dir = Path(__file__).resolve().parent
    if str(scripts_dir) not in sys.path:
        sys.path.insert(0, str(scripts_dir))

    if args.fallback_docx:
        out = build_document_docx(out_path)
    else:
        try:
            out = build_document_com(out_path)
        except Exception as exc:
            print("Word COM failed:", exc)
            print("Trying python-docx fallback...")
            out = build_document_docx(out_path)

    print("Created:", out.as_posix())


if __name__ == "__main__":
    main()
